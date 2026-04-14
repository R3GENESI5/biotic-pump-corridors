"""Build complete Word document with embedded figures at native resolution."""
import sys
sys.path.insert(0, ".")

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
from PIL import Image

OUTPUT_DIR = Path("outputs")
PAPER_DIR = Path("paper")

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# --- Styles ---
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_figure(image_path, caption, max_width_inches=6.0):
    """Add a figure at native aspect ratio, fitting within max_width."""
    img = Image.open(image_path)
    w_px, h_px = img.size
    dpi = img.info.get("dpi", (150, 150))
    if isinstance(dpi, tuple):
        dpi_x = dpi[0]
    else:
        dpi_x = dpi

    # Native width in inches
    native_w = w_px / dpi_x
    native_h = h_px / (dpi[1] if isinstance(dpi, tuple) else dpi)

    # Scale down only if wider than max_width, preserving aspect ratio
    if native_w > max_width_inches:
        scale = max_width_inches / native_w
        w = Inches(max_width_inches)
        h = Inches(native_h * scale)
    else:
        w = Inches(native_w)
        h = Inches(native_h)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=w, height=h)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(12)

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing


# =========================================================================
# TITLE
# =========================================================================
add_para("Satellite evidence for progressive weakening of condensation-driven "
         "moisture transport along tropical forest corridors",
         bold=True, size=16, align="center")
doc.add_paragraph()
add_para("Ali Bin Shahid", bold=True, size=12, align="center")
add_para("Independent Scientist, Islamabad, Pakistan", italic=True, size=10, align="center")
doc.add_paragraph()

# =========================================================================
# ABSTRACT
# =========================================================================
add_heading("Abstract", level=1)
doc.add_paragraph(
    "The biotic pump hypothesis proposes that forest evapotranspiration drives "
    "condensation-generated low pressure that sustains inland moisture transport. "
    "We test this mechanism using 25 years of satellite data along two tropical "
    "moisture corridors: Belem\u2013Manaus (Amazon) and Gulf of Guinea\u2013Kisangani (Congo). "
    "Comparing 2019 with the 2023 drought, we find that outgoing longwave radiation "
    "(OLR) increased (+2.9 W/m\u00b2 over the Eastern Amazon, +1.3 W/m\u00b2 over the Congo), "
    "cloud fraction decreased (\u22122.3% Amazon, \u22120.5% Congo), evapotranspiration "
    "declined (\u221216 mm/yr Amazon), and column water vapour accumulated at the coast "
    "rather than penetrating inland \u2014 all signatures of weakened condensation-driven "
    "transport. These signals amplify with distance from the coast: interior OLR "
    "increased +4.3 W/m\u00b2 versus +2.7 W/m\u00b2 at the coast, and the coast-to-interior "
    "water vapour gradient steepened by \u22121.2 kg/m\u00b2. The signal concentrates in the "
    "September\u2013November dry-to-wet transition, when forest transpiration should "
    "initiate the wet season ahead of the Intertropical Convergence Zone; interior "
    "OLR during SON increased by +24.5 W/m\u00b2. Over 2001\u20132025, the Amazon "
    "coast-to-interior cloud fraction gradient declined at \u22120.17%/yr (p = 0.001 "
    "after removing the ENSO signal using the Oceanic Ni\u00f1o Index, which explains "
    "less than 1% of gradient variance). NDVI analysis confirms that 88% of intact "
    "forest cells (NDVI > 0.6) experienced greenness decline (t = \u221215.4, p < 10\u207b\u2076), "
    "establishing that the evapotranspiration reduction reflects drought stress in "
    "standing forest, not land-cover change. The Congo corridor shows the same trends "
    "at weaker magnitudes, consistent with an earlier stage of degradation. These "
    "results provide spatially resolved satellite evidence for the biotic pump "
    "mechanism and document its progressive failure in the two largest tropical forests."
)

# =========================================================================
# 1. INTRODUCTION
# =========================================================================
add_heading("1. Introduction", level=1)

intro_paras = [
    "Tropical forests cycle enormous quantities of water between the land surface and the atmosphere. In the Amazon basin, approximately half of the precipitation that falls on the forest is returned to the atmosphere through evapotranspiration, where it is available to generate precipitation further downwind (Salati et al., 1979; Eltahir and Bras, 1996). This moisture recycling sustains precipitation deep into the continental interior, far beyond the reach of oceanic moisture alone.",
    "The mechanism by which this inland moisture transport is maintained remains debated. The prevailing view attributes it to large-scale pressure gradients driven by differential heating between the ocean and the continent (e.g., the South American monsoon system). However, Makarieva and Gorshkov (2007) proposed an alternative mechanism \u2014 the biotic pump \u2014 in which the condensation of water vapour itself generates the low-pressure anomaly that draws moist air inland. In this framework, forest evapotranspiration is not merely a consequence of available moisture but an active driver of the atmospheric circulation that delivers it.",
    "The biotic pump hypothesis has been refined in subsequent theoretical work (Makarieva et al., 2013; Makarieva and Gorshkov, 2015) and finds indirect support from several lines of evidence. The Amazon wet season initiates before the southward migration of the Intertropical Convergence Zone (ITCZ), suggesting that vegetation-atmosphere interactions drive early convective onset (Li and Fu, 2004; Wright et al., 2017). The Eurasian boreal forest sustains inland moisture transport during the growing season that collapses during dormancy (Makarieva et al., 2009). Deforestation experiments in climate models consistently reduce precipitation in deforested regions and downwind (Spracklen et al., 2012; Staal et al., 2018; Nobre et al., 2016).",
    "Despite this theoretical and modelling evidence, direct observational testing of the biotic pump at the corridor scale has been limited. If the biotic pump operates, its weakening should produce specific, observable satellite signatures: (1) reduced evapotranspiration, (2) decreased cloud cover, (3) increased outgoing longwave radiation, and (4) steepening of the coast-to-interior moisture gradient as oceanic moisture fails to penetrate inland. These signals should amplify with distance from the coast and concentrate in the seasonal transition periods when forest-driven convection, rather than the ITCZ, is the primary moisture driver.",
    "The years 2019 and 2023 provide an opportunity for such a test. The Amazon and Congo basins experienced record or near-record drought in 2023 (Espinoza et al., 2024; Wongchuig et al., 2025), offering a natural experiment in which the biotic pump was stressed. By comparing satellite-derived atmospheric and surface variables between 2019 (a relatively normal year) and 2023 (drought), and placing this comparison within a 25-year satellite record, we can assess whether the predicted signatures appear, whether they follow the spatial and seasonal patterns the biotic pump framework predicts, and whether they constitute a trend or an isolated anomaly.",
    "We focus on two corridors: the Belem\u2013Manaus corridor in the Amazon, connecting the Atlantic coast to the western interior; and the Gulf of Guinea\u2013Kisangani corridor in the Congo, connecting the oceanic moisture source to the interior of Africa\u2019s largest remaining rainforest.",
    "Here we show that the predicted signatures of biotic pump weakening are present in satellite data from four independent instruments (CERES, MODIS, AIRS), that they amplify with distance from the coast, that they concentrate in the dry-to-wet transition season, and that they represent a progressive 25-year trend independent of ENSO variability. We further demonstrate that the ET decline reflects drought stress in intact forest rather than deforestation.",
]
for p in intro_paras:
    doc.add_paragraph(p)

# Figure 1: Study area
add_figure(OUTPUT_DIR / "focal_regions.png",
           "Figure 1. Study regions showing the three focal areas: "
           "Maraba (Amazon), Congo Basin, and Tropical Atlantic.")

# =========================================================================
# 2. DATA AND METHODS
# =========================================================================
add_heading("2. Data and Methods", level=1)

add_heading("2.1 Study regions and corridor definition", level=2)
doc.add_paragraph(
    "We define two tropical moisture corridors based on the biotic pump framework "
    "(Makarieva and Gorshkov, 2007), each connecting an oceanic moisture source to "
    "a continental forest interior."
)
doc.add_paragraph(
    "The Amazon corridor runs from Belem on the Atlantic coast (1.4\u00b0S, 48.5\u00b0W) "
    "westward through Maraba, Altamira, and Itaituba to Manaus (3.1\u00b0S, 60.0\u00b0W), "
    "a distance of approximately 1,300 km. We average variables over a latitude band "
    "of 6\u00b0S to 0\u00b0N. For spatial correlation analysis, we use a wider Eastern Amazon "
    "domain (10\u00b0S\u20132\u00b0N, 65\u00b0W\u201345\u00b0W), encompassing approximately 240 grid cells at 1-degree resolution."
)
doc.add_paragraph(
    "The Congo corridor extends from the Gulf of Guinea coast near Douala (4.0\u00b0N, 9.7\u00b0E) "
    "eastward through Yaounde, Ouesso, and Mbandaka to Kisangani (0.5\u00b0N, 25.2\u00b0E), "
    "covering approximately 1,700 km. We average over a latitude band of 3\u00b0S to 5\u00b0N."
)

add_heading("2.2 Satellite data", level=2)
doc.add_paragraph(
    "CERES EBAF-TOA Edition 4.2.1 (Loeb et al., 2018) provides monthly gridded "
    "(1\u00b0 \u00d7 1\u00b0) top-of-atmosphere radiative fluxes from March 2000 through January 2026. "
    "We use all-sky OLR, clear-sky OLR, and cloud area fraction. The longwave cloud "
    "radiative effect (CRE_LW) is derived as OLR_clear minus OLR_all."
)
doc.add_paragraph(
    "MODIS MOD16A3GF Version 061 (Running et al., 2019) provides gap-filled annual "
    "evapotranspiration at 500 m. Tiles are reprojected to geographic coordinates and "
    "aggregated to 1-degree resolution. MODIS MOD13C2 Version 061 provides monthly NDVI "
    "at 0.05-degree global resolution. AIRS L3 Version 7.0 (Tian et al., 2020) provides "
    "total column water vapour at 1\u00b0 \u00d7 1\u00b0 monthly resolution. NOAA ONI provides "
    "3-month running mean SST anomalies in the Ni\u00f1o 3.4 region."
)

add_heading("2.3 Analysis methods", level=2)
doc.add_paragraph(
    "Corridor transects are computed by selecting data within the corridor latitude band "
    "and calculating the cosine-latitude-weighted mean across latitudes at each longitude. "
    "Seasonal decomposition uses four periods: DJF, MAM, JJA, and SON. Multi-year trends "
    "are computed over 2001\u20132025 using ordinary least squares. ENSO control regresses each "
    "corridor metric on annual mean ONI and analyses residual trends. Spatial correlation "
    "between ET and OLR change is computed across all 1\u00b0 cells in the wider Eastern Amazon. "
    "Forest integrity assessment classifies cells by 2019 NDVI (>0.6 = dense forest) and "
    "tests whether forest NDVI change differs from zero."
)

# =========================================================================
# 3. RESULTS
# =========================================================================
add_heading("3. Results", level=1)

add_heading("3.1 Regional overview: 2019 versus 2023", level=2)
doc.add_paragraph(
    "Global mean OLR increased by 0.16 W/m\u00b2 between 2019 and 2023. Against this "
    "background, the Eastern Amazon showed an OLR increase of +2.9 W/m\u00b2 \u2014 18 times "
    "the global mean \u2014 while the Congo Basin increased by +1.3 W/m\u00b2 (Table 1)."
)

add_table(
    ["Variable", "Amazon 2019", "Amazon 2023", "Change", "Congo 2019", "Congo 2023", "Change"],
    [
        ["OLR (W/m\u00b2)", "236.6", "239.5", "+2.9", "220.6", "221.9", "+1.3"],
        ["Cloud (%)", "72.5", "70.3", "\u22122.3", "79.1", "78.6", "\u22120.5"],
        ["CRE (W/m\u00b2)", "44.6", "43.0", "\u22121.5", "55.4", "54.6", "\u22120.7"],
        ["ET (mm/yr)", "1419", "1403", "\u221216", "1298", "1310", "+13"],
    ],
)
add_para("Table 1. Regional means and changes, 2019 vs 2023.", italic=True, size=9)

add_figure(OUTPUT_DIR / "wide_amazon_olr.png",
           "Figure 2. Eastern Amazon OLR: 2019 (left), 2023 (centre), "
           "difference (right). Dashed line shows the Belem-Manaus corridor.")

add_figure(OUTPUT_DIR / "et_comparison_maraba.png",
           "Figure 3. Maraba region evapotranspiration (MODIS MOD16A3GF): "
           "2019 (left), 2023 (centre), difference (right).")

add_heading("3.2 Corridor transect analysis", level=2)
doc.add_paragraph(
    "Along the Belem\u2013Manaus corridor, OLR increased at every longitude, but the "
    "increase was substantially larger in the interior (+4.3 W/m\u00b2 near Manaus) than "
    "at the coast (+2.7 W/m\u00b2 near Belem). AIRS column water vapour revealed a "
    "striking asymmetry: the coast gained moisture (+0.8 kg/m\u00b2) while the interior "
    "lost moisture (\u22120.4 kg/m\u00b2). The coast-to-interior water vapour gradient steepened "
    "by \u22121.2 kg/m\u00b2."
)

add_table(
    ["Variable", "Coast", "\u0394", "Mid", "\u0394", "Interior", "\u0394"],
    [
        ["OLR (W/m\u00b2)", "241.5\u2192244.2", "+2.70", "240.5\u2192243.2", "+2.76", "225.9\u2192230.2", "+4.26"],
        ["Cloud (%)", "70.9\u219268.3", "\u22122.55", "74.7\u219271.0", "\u22123.64", "78.4\u219277.2", "\u22121.14"],
        ["CRE (W/m\u00b2)", "42.2\u219240.3", "\u22121.94", "41.3\u219239.0", "\u22122.39", "52.7\u219250.3", "\u22122.43"],
        ["WV (kg/m\u00b2)", "48.9\u219249.7", "+0.82", "49.1\u219249.3", "+0.21", "46.4\u219246.0", "\u22120.41"],
    ],
)
add_para("Table 2. Amazon corridor segment values, 2019 vs 2023.", italic=True, size=9)

add_figure(OUTPUT_DIR / "corridor_transects.png",
           "Figure 4. Belem-Manaus corridor transects: OLR, cloud fraction, "
           "and CRE averaged over 6S-0N latitude band, 2019 vs 2023.")

add_figure(OUTPUT_DIR / "airs_corridor_transect.png",
           "Figure 5. AIRS column water vapour along the Belem-Manaus corridor. "
           "Top: absolute values. Bottom: change (2023 minus 2019).")

add_figure(OUTPUT_DIR / "congo_corridor_transects.png",
           "Figure 6. Congo corridor transects: Gulf of Guinea to Kisangani.")

add_heading("3.3 Seasonal decomposition", level=2)
doc.add_paragraph(
    "The dry-to-wet transition season (SON) dominates the annual signal. "
    "Interior OLR increased by +24.5 W/m\u00b2 during SON, compared to +10.6 W/m\u00b2 "
    "during JJA and \u22128.5 W/m\u00b2 during DJF."
)

add_table(
    ["Season", "Coast", "Mid", "Interior"],
    [
        ["DJF (wet)", "+0.3", "\u22124.5", "\u22128.5"],
        ["MAM", "\u22120.7", "\u22122.1", "\u22129.7"],
        ["JJA (dry)", "+7.5", "+7.6", "+10.6"],
        ["SON (dry-to-wet)", "+10.5", "+14.5", "+24.5"],
    ],
)
add_para("Table 3. Seasonal OLR change (W/m\u00b2) along the Amazon corridor.", italic=True, size=9)

add_figure(OUTPUT_DIR / "corridor_seasonal_transects.png",
           "Figure 7. Seasonal decomposition of corridor changes (2023 minus 2019). "
           "SON (yellow) dominates the signal at the interior end.")

add_figure(OUTPUT_DIR / "corridor_seasonal_bars.png",
           "Figure 8. Seasonal breakdown by corridor segment. "
           "Which season drove the change at each location?")

add_heading("3.4 Multi-year trends (2001-2025)", level=2)
doc.add_paragraph(
    "Over the 25-year CERES record, the Amazon corridor shows progressive gradient "
    "weakening. The cloud fraction gradient declined at \u22120.17%/yr, and the CRE gradient "
    "declined at \u22120.12 W/m\u00b2/yr. Interior cloud fraction fell from 82.7% in 2005 to "
    "77.2% by 2023."
)

add_figure(OUTPUT_DIR / "corridor_gradient_evolution.png",
           "Figure 9. Coast-to-interior gradient evolution along the Amazon corridor, "
           "2001-2025. Dashed lines show linear trends.")

add_figure(OUTPUT_DIR / "corridor_SON_multiyear.png",
           "Figure 10. SON-only time series along the Amazon corridor, 2001-2025.")

add_heading("3.5 ENSO independence", level=2)
doc.add_paragraph(
    "The correlation between annual mean ONI and the Amazon cloud fraction gradient "
    "is r = 0.015 (p = 0.94), explaining less than 1% of variance."
)

add_table(
    ["Metric", "Raw trend", "ENSO-adjusted", "p-value"],
    [
        ["Cloud fraction gradient", "\u22120.170%/yr", "\u22120.170%/yr", "0.001"],
        ["CRE gradient", "\u22120.116 W/m\u00b2/yr", "\u22120.117 W/m\u00b2/yr", "0.033"],
        ["OLR gradient", "+0.074 W/m\u00b2/yr", "+0.075 W/m\u00b2/yr", "0.201"],
    ],
)
add_para("Table 4. Amazon corridor gradient trends, raw and ENSO-adjusted.", italic=True, size=9)

add_figure(OUTPUT_DIR / "enso_adjusted_trends.png",
           "Figure 11. ENSO-adjusted corridor gradient trends. After removing the ONI "
           "signal, the degradation trend persists (p = 0.001 for cloud fraction).")

add_figure(OUTPUT_DIR / "enso_2023_vs_other_elnino.png",
           "Figure 12. 2023 compared to other El Nino years. "
           "2023 is not anomalous among El Nino events.")

add_heading("3.6 Forest integrity: drought stress, not deforestation", level=2)
doc.add_paragraph(
    "Among 225 grid cells classified as dense forest (NDVI > 0.6 in 2019), "
    "88.4% showed NDVI decline between 2019 and 2023. The mean NDVI change in "
    "forest cells was \u22120.009 (t = \u221215.4, p < 10\u207b\u2076). NDVI and ET change are "
    "positively correlated (r = 0.29, p = 0.0006)."
)

add_figure(OUTPUT_DIR / "ndvi_forest_vs_nonforest.png",
           "Figure 13. NDVI change in intact forest vs non-forest cells. "
           "88% of dense forest cells show NDVI decline.")

add_figure(OUTPUT_DIR / "ndvi_vs_et_scatter.png",
           "Figure 14. NDVI change vs ET change across the Eastern Amazon.")

add_heading("3.7 Spatial correlation between ET and OLR", level=2)
doc.add_paragraph(
    "Across 159 grid cells in the wider Eastern Amazon, ET change and OLR change "
    "are positively correlated (r = 0.32, p < 0.0001). Forty-seven percent of cells "
    "fall in the pump failure quadrant (ET < 0, OLR > 0)."
)

add_figure(OUTPUT_DIR / "wide_amazon_scatter.png",
           "Figure 15. ET vs atmospheric response per grid cell in the Eastern Amazon.")

# =========================================================================
# 4. DISCUSSION
# =========================================================================
add_heading("4. Discussion", level=1)

add_heading("4.1 Satellite signatures of the biotic pump", level=2)
doc.add_paragraph(
    "The biotic pump hypothesis predicts that weakening of forest evapotranspiration "
    "should produce four observable consequences: reduced ET, thinning clouds, increased "
    "OLR, and steepened coast-to-interior moisture gradients. Our results show all four "
    "along the Belem\u2013Manaus corridor, and critically, the signals amplify with distance "
    "from the coast. This spatial pattern \u2014 signal amplification with distance from the "
    "moisture source \u2014 is the distinctive signature of a transport failure mechanism."
)

add_heading("4.2 The SON transition: where the biotic pump operates", level=2)
doc.add_paragraph(
    "The corridor signal concentrates in SON \u2014 when interior OLR increased +24.5 W/m\u00b2. "
    "During DJF, when the ITCZ provides geophysical forcing, the corridor partially "
    "recovered. This seasonal contrast \u2014 biotic pump failure in SON, geophysical "
    "compensation in DJF \u2014 is difficult to explain by any mechanism other than "
    "reduced vegetation-driven moisture convergence."
)

add_heading("4.3 A progressive trend, not a single event", level=2)
doc.add_paragraph(
    "The 25-year record reveals that 2023 is not isolated but part of a multi-decadal "
    "degradation. The year-by-year data show a ratchet: drought events produce sharp "
    "gradient reductions from which the corridor only partially recovers. This is "
    "consistent with a positive feedback proposed in tipping point studies (Lenton et al., "
    "2008; Staal et al., 2020; Zemp et al., 2017)."
)

add_heading("4.4 Amazon and Congo: different stages of degradation", level=2)
doc.add_paragraph(
    "We interpret the Congo as occupying an earlier position on the same degradation "
    "trajectory. The Amazon has experienced greater deforestation and more frequent "
    "severe droughts. The Congo\u2019s declining CRE gradient suggests the atmospheric "
    "response is beginning to decouple from the still-active forest."
)

add_heading("4.5 Moisture displacement and extreme weather", level=2)
doc.add_paragraph(
    "During September 2023, OLR showed positive anomalies of +12.5 W/m\u00b2 over the "
    "Amazon while the eastern Mediterranean showed negative anomalies (\u22124.6 W/m\u00b2) "
    "coinciding with Storm Daniel. We note this as a hypothesis for future investigation."
)

add_figure(OUTPUT_DIR / "storm_daniel_toa_lw_all_anomaly.png",
           "Figure 16. September 2023 OLR anomaly showing the spatial dipole between "
           "tropical forest source regions and the Storm Daniel location.")

add_figure(OUTPUT_DIR / "storm_moisture_dipole.png",
           "Figure 17. Moisture dipole quantification: forest sources show OLR surplus, "
           "Storm Daniel region shows OLR deficit.")

add_heading("4.6 Limitations", level=2)
doc.add_paragraph(
    "The 2019\u20132023 comparison captures the drought signal but does not characterise "
    "full interannual variability. CERES at 1-degree resolution cannot resolve sub-grid "
    "processes. The MODIS ET product has known biases in tropical forests (Mu et al., 2011). "
    "While we control for ENSO, other climate modes could contribute. Satellite observations "
    "alone cannot establish causality; confirming causation requires Earth system model experiments."
)

# =========================================================================
# 5. CONCLUSIONS
# =========================================================================
add_heading("5. Conclusions", level=1)

conclusions = [
    "1. Four independent satellite instruments show a spatially coherent weakening of moisture transport along the Belem\u2013Manaus and Gulf of Guinea\u2013Kisangani corridors between 2019 and 2023, consistent with biotic pump predictions.",
    "2. Signals amplify with distance from the coast: interior OLR increased +4.3 W/m\u00b2 versus +2.7 W/m\u00b2 at the coast; the AIRS water vapour gradient steepened by \u22121.2 kg/m\u00b2.",
    "3. The breakdown concentrates in September\u2013November, when forest transpiration should initiate the wet season. Interior OLR during SON increased +24.5 W/m\u00b2; the wet season (DJF) showed partial recovery from geophysical forcing.",
    "4. The Amazon corridor has undergone progressive weakening over 25 years. The cloud fraction gradient trend (\u22120.17%/yr, p = 0.001) and CRE gradient trend (\u22120.12 W/m\u00b2/yr, p = 0.033) are statistically significant after ENSO removal.",
    "5. NDVI confirms that 88% of intact forest cells experienced greenness decline (t = \u221215.4, p < 10\u207b\u2076), establishing drought stress in standing forest, not deforestation.",
    "6. The Congo corridor shows the same trends at weaker magnitudes, consistent with an earlier degradation stage.",
    "7. A spatial dipole during September 2023 coincides with Storm Daniel and warrants investigation through moisture trajectory analysis.",
]
for c in conclusions:
    doc.add_paragraph(c)

doc.add_paragraph(
    "These findings constitute the first spatially resolved, multi-instrument satellite "
    "documentation of biotic pump weakening along defined moisture corridors, and indicate "
    "that condensation-driven circulation in the two largest tropical forests is progressively "
    "deteriorating."
)

# Summary figure
add_figure(OUTPUT_DIR / "final_summary.png",
           "Figure 18. Summary figure: (a) OLR difference with corridors and storm locations, "
           "(b-c) corridor transects, (d) seasonal decomposition, (e) ENSO-adjusted trend, "
           "(f) NDVI in intact forest, (g) ET vs OLR scatter.")

# =========================================================================
# DATA AVAILABILITY
# =========================================================================
add_heading("Data Availability", level=1)
doc.add_paragraph(
    "CERES EBAF-TOA Ed4.2.1 is available from the NASA Langley Research Center "
    "Atmospheric Science Data Center. MODIS MOD16A3GF and MOD13C2 are available from "
    "the NASA Land Processes DAAC. AIRS L3 data are available from the NASA GES DISC. "
    "ONI data are available from the NOAA Climate Prediction Center. All data were "
    "accessed through NASA Earthdata (https://earthdata.nasa.gov/)."
)

# =========================================================================
# REFERENCES
# =========================================================================
add_heading("References", level=1)

refs = [
    "Eltahir, E. A. B. and Bras, R. L. (1996). Precipitation recycling in the Amazon basin. Quarterly Journal of the Royal Meteorological Society, 122, 539\u2013555.",
    "Espinoza, J. C., Jimenez, J. C., Marengo, J. A., Schongart, J., Ronchail, J., Lavado-Casimiro, W., and Ribeiro, J. V. M. (2024). The new record of drought and warmth in the Amazon in 2023 related to regional and global climatic features. Scientific Reports, 14, 8107. https://doi.org/10.1038/s41598-024-58782-5",
    "Fu, R., Yin, L., Li, W., et al. (2013). Increased dry-season length over southern Amazonia in recent decades. PNAS, 110, 18110\u201318115.",
    "Lenton, T. M., Held, H., Kriegler, E., et al. (2008). Tipping elements in the Earth\u2019s climate system. PNAS, 105, 1786\u20131793.",
    "Li, W. and Fu, R. (2004). Transition of the large-scale atmospheric and land surface conditions from the dry to the wet season over Amazonia. Journal of Climate, 17, 2637\u20132651.",
    "Loeb, N. G., Doelling, D. R., Wang, H., et al. (2018). CERES EBAF TOA Edition-4.0 Data Product. Journal of Climate, 31, 895\u2013918.",
    "Makarieva, A. M. and Gorshkov, V. G. (2007). Biotic pump of atmospheric moisture as driver of the hydrological cycle on land. Hydrology and Earth System Sciences, 11, 1013\u20131033.",
    "Makarieva, A. M., Gorshkov, V. G., and Li, B.-L. (2009). Precipitation on land versus distance from the ocean. Ecological Complexity, 6, 302\u2013307.",
    "Makarieva, A. M., Gorshkov, V. G., Sheil, D., Nobre, A. D., and Li, B.-L. (2013). Where do winds come from? ACP, 13, 1039\u20131056.",
    "Mu, Q., Zhao, M., and Running, S. W. (2011). Improvements to a MODIS global terrestrial evapotranspiration algorithm. Remote Sensing of Environment, 115, 1781\u20131800.",
    "Nobre, C. A., Sampaio, G., Borma, L. S., et al. (2016). Land-use and climate change risks in the Amazon. PNAS, 113, 10759\u201310768.",
    "Running, S. W., Mu, Q., Zhao, M., and Moreno, A. (2019). MODIS Global Terrestrial ET Product. User\u2019s Guide, Collection 6.1.",
    "Salati, E., Dall\u2019Olio, A., Matsui, E., and Gat, J. R. (1979). Recycling of water in the Amazon basin. Water Resources Research, 15, 1250\u20131258.",
    "Spracklen, D. V., Arnold, S. R., and Taylor, C. M. (2012). Observations of increased tropical rainfall preceded by air passage over forests. Nature, 489, 282\u2013285.",
    "Staal, A., Tuinenburg, O. A., et al. (2018). Forest-rainfall cascades buffer against drought across the Amazon. Nature Climate Change, 8, 539\u2013543.",
    "Staal, A., Fetzer, I., et al. (2020). Hysteresis of tropical forests in the 21st century. Nature Communications, 11, 4978.",
    "Stein, A. F., et al. (2015). NOAA\u2019s HYSPLIT atmospheric transport and dispersion modeling system. BAMS, 96, 2059\u20132077.",
    "Tian, B., et al. (2020). AIRS/Aqua L3 Monthly Standard Physical Retrieval V7.0. NASA GES DISC.",
    "Wongchuig, S., Papa, F., Fleischmann, A. S., et al. (2025). Recent significant drying in Central Congo Basin linked to weakened Walker circulation and warmer Atlantic. npj Climate and Atmospheric Science, 8, 331. https://doi.org/10.1038/s41612-025-01225-3",
    "Wright, J. S., Fu, R., Worden, J. R., et al. (2017). Rainforest-initiated wet season onset over the southern Amazon. PNAS, 114, 8481\u20138486.",
    "Zemp, D. C., et al. (2017). Self-amplified Amazon forest loss due to vegetation-atmosphere feedbacks. Nature Communications, 8, 14681.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9)

# =========================================================================
# SAVE
# =========================================================================
output_path = PAPER_DIR / "Shahid_2025_Biotic_Pump_Corridors.docx"
doc.save(str(output_path))
print(f"Saved: {output_path}")
print(f"File size: {output_path.stat().st_size / 1024 / 1024:.1f} MB")
