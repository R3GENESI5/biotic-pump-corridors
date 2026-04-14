"""Build supplementary materials Word document with all supporting figures."""
import sys
sys.path.insert(0, ".")

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from PIL import Image

OUTPUT_DIR = Path("outputs")
PAPER_DIR = Path("paper")

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_figure(image_path, caption, max_width_inches=6.0):
    if not image_path.exists():
        add_para(f"[Missing: {image_path.name}]", italic=True, size=9)
        return

    img = Image.open(image_path)
    w_px, h_px = img.size
    dpi = img.info.get("dpi", (150, 150))
    dpi_x = dpi[0] if isinstance(dpi, tuple) else dpi
    dpi_y = dpi[1] if isinstance(dpi, tuple) else dpi

    native_w = w_px / dpi_x
    native_h = h_px / dpi_y

    if native_w > max_width_inches:
        scale = max_width_inches / native_w
        w = Inches(max_width_inches)
        h = Inches(native_h * scale)
    else:
        w = Inches(native_w)
        h = Inches(native_h)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=w, height=h)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(12)


# =========================================================================
# TITLE
# =========================================================================
add_para("Supplementary Materials", bold=True, size=16, align="center")
doc.add_paragraph()
add_para("Satellite evidence for progressive weakening of condensation-driven "
         "moisture transport along tropical forest corridors",
         italic=True, size=12, align="center")
doc.add_paragraph()
add_para("Ali Bin Shahid", bold=True, size=11, align="center")
add_para("Independent Scientist, Islamabad, Pakistan", italic=True, size=10, align="center")
doc.add_paragraph()
doc.add_paragraph()

# =========================================================================
# S1. STORM DISPLACEMENT ANALYSIS
# =========================================================================
add_heading("S1. Storm Displacement Analysis", level=1)

doc.add_paragraph(
    "These figures show the September OLR and cloud fraction anomalies "
    "(relative to the 2001\u20132018 September climatology) during Storm Daniel "
    "(September 2023) and Storm Boris (September 2024). Forest source regions "
    "(Amazon, Congo) and storm locations are marked."
)

add_figure(OUTPUT_DIR / "storm_daniel_toa_lw_all_anomaly.png",
           "Figure S1. September 2023 OLR anomaly with forest source regions "
           "and Storm Daniel location (Libya). Red = positive OLR anomaly "
           "(drier/less cloud), blue = negative (wetter/more convection).")

add_figure(OUTPUT_DIR / "storm_daniel_cldarea_anomaly.png",
           "Figure S2. September 2023 cloud fraction anomaly. Brown = fewer "
           "clouds than normal, green = more clouds than normal.")

add_figure(OUTPUT_DIR / "storm_daniel_cre_lw_anomaly.png",
           "Figure S3. September 2023 longwave cloud radiative effect anomaly.")

add_figure(OUTPUT_DIR / "storm_daniel_airs_wv_anomaly.png",
           "Figure S4. September 2023 AIRS column water vapour anomaly "
           "(relative to September 2019). Brown = drier, green = wetter.")

add_figure(OUTPUT_DIR / "storm_boris_toa_lw_all_anomaly.png",
           "Figure S5. September 2024 OLR anomaly with Storm Boris location "
           "(Central Europe).")

add_figure(OUTPUT_DIR / "storm_boris_cldarea_anomaly.png",
           "Figure S6. September 2024 cloud fraction anomaly with Storm Boris.")

add_figure(OUTPUT_DIR / "storm_moisture_dipole.png",
           "Figure S7. Moisture dipole quantification for September 2023. "
           "Left: OLR anomaly by region. Right: cloud fraction anomaly. "
           "Forest sources show surplus OLR, Storm Daniel region shows deficit.")

# =========================================================================
# S2. FULL ENSO CONTROL
# =========================================================================
add_heading("S2. Full ENSO Control Analysis", level=1)

doc.add_paragraph(
    "These figures show the complete ENSO control analysis for both corridors, "
    "including all three variables (OLR, cloud fraction, CRE)."
)

add_figure(OUTPUT_DIR / "enso_amazon_gradient.png",
           "Figure S8. Amazon corridor gradient evolution coloured by ENSO phase "
           "(red = El Nino, blue = La Nina, grey = neutral). Dashed lines show "
           "trends within each phase.")

add_figure(OUTPUT_DIR / "enso_congo_gradient.png",
           "Figure S9. Congo corridor gradient evolution coloured by ENSO phase.")

# =========================================================================
# S3. MULTI-YEAR TIME SERIES
# =========================================================================
add_heading("S3. Multi-Year Time Series", level=1)

add_figure(OUTPUT_DIR / "corridor_multiyear_timeseries.png",
           "Figure S10. Amazon corridor annual mean time series (2001\u20132025) "
           "for OLR, cloud fraction, and CRE at coast, mid-corridor, and interior segments.")

add_figure(OUTPUT_DIR / "congo_corridor_multiyear.png",
           "Figure S11. Congo corridor annual mean time series (2001\u20132025).")

add_figure(OUTPUT_DIR / "congo_corridor_gradient_evolution.png",
           "Figure S12. Congo corridor coast-to-interior gradient evolution (2001\u20132025).")

# =========================================================================
# S4. MONTHLY TIME SERIES
# =========================================================================
add_heading("S4. Monthly Time Series", level=1)

add_figure(OUTPUT_DIR / "olr_timeseries.png",
           "Figure S13. Monthly OLR time series (2000\u20132025) for all three "
           "focal regions with 12-month rolling mean. Pink shading = focus years.")

add_figure(OUTPUT_DIR / "cloud_fraction_timeseries.png",
           "Figure S14. Monthly cloud fraction time series (2000\u20132025).")

# =========================================================================
# S5. CONGO CORRIDOR DETAILS
# =========================================================================
add_heading("S5. Congo Corridor Details", level=1)

add_figure(OUTPUT_DIR / "congo_corridor_map_toa_lw_all.png",
           "Figure S15. Congo corridor OLR spatial maps: 2019, 2023, and difference. "
           "Dashed line shows the Douala\u2013Kisangani corridor path.")

add_figure(OUTPUT_DIR / "congo_corridor_map_cldarea.png",
           "Figure S16. Congo corridor cloud fraction spatial maps.")

add_figure(OUTPUT_DIR / "congo_corridor_map_cre_lw.png",
           "Figure S17. Congo corridor CRE spatial maps.")

add_figure(OUTPUT_DIR / "congo_corridor_gradient_bars.png",
           "Figure S18. Congo corridor: change vs distance from coast.")

add_figure(OUTPUT_DIR / "congo_corridor_seasonal.png",
           "Figure S19. Congo corridor seasonal decomposition of changes (2023\u20132019).")

# =========================================================================
# S6. AMAZON CORRIDOR DETAILS
# =========================================================================
add_heading("S6. Amazon Corridor Details", level=1)

add_figure(OUTPUT_DIR / "corridor_map_toa_lw_all.png",
           "Figure S20. Amazon corridor OLR spatial maps with city markers "
           "and corridor path overlay.")

add_figure(OUTPUT_DIR / "corridor_map_cldarea.png",
           "Figure S21. Amazon corridor cloud fraction spatial maps.")

add_figure(OUTPUT_DIR / "corridor_map_cre_lw.png",
           "Figure S22. Amazon corridor CRE spatial maps.")

add_figure(OUTPUT_DIR / "corridor_gradient_bars.png",
           "Figure S23. Amazon corridor: change vs distance from Atlantic coast (km).")

add_figure(OUTPUT_DIR / "olr_coast_interior_amazon.png",
           "Figure S24. Coast vs interior OLR monthly time series for the Amazon, "
           "testing Makarieva\u2019s prediction that the coastal strip is critical.")

add_figure(OUTPUT_DIR / "airs_corridor_map.png",
           "Figure S25. AIRS water vapour spatial maps along the Amazon corridor.")

# =========================================================================
# S7. REGIONAL COMPARISON MAPS
# =========================================================================
add_heading("S7. Regional Comparison Maps", level=1)

add_figure(OUTPUT_DIR / "olr_comparison_maraba.png",
           "Figure S26. Maraba subregion OLR comparison: 2019, 2023, difference.")

add_figure(OUTPUT_DIR / "cloud_fraction_maraba.png",
           "Figure S27. Maraba subregion cloud fraction comparison.")

add_figure(OUTPUT_DIR / "cre_comparison_maraba.png",
           "Figure S28. Maraba subregion CRE comparison.")

add_figure(OUTPUT_DIR / "olr_comparison_congo.png",
           "Figure S29. Congo Basin OLR comparison.")

add_figure(OUTPUT_DIR / "cloud_fraction_congo.png",
           "Figure S30. Congo Basin cloud fraction comparison.")

add_figure(OUTPUT_DIR / "cre_comparison_congo.png",
           "Figure S31. Congo Basin CRE comparison.")

add_figure(OUTPUT_DIR / "olr_tropical_overview.png",
           "Figure S32. Tropical overview (30S\u201330N): OLR 2019, 2023, and difference.")

# =========================================================================
# S8. WIDER AMAZON ANALYSIS
# =========================================================================
add_heading("S8. Wider Amazon Analysis", level=1)

add_figure(OUTPUT_DIR / "wide_amazon_cloud_fraction.png",
           "Figure S33. Eastern Amazon (10S\u20132N, 65W\u201345W) cloud fraction: "
           "2019, 2023, difference.")

add_figure(OUTPUT_DIR / "wide_amazon_lw_cloud_radiative_effect.png",
           "Figure S34. Eastern Amazon CRE: 2019, 2023, difference.")

# =========================================================================
# S9. NDVI AND ET DETAILS
# =========================================================================
add_heading("S9. NDVI and ET Details", level=1)

add_figure(OUTPUT_DIR / "ndvi_eastern_amazon.png",
           "Figure S35. Eastern Amazon NDVI comparison: 2019, 2023, difference. "
           "At 0.25-degree resolution.")

add_figure(OUTPUT_DIR / "ndvi_congo_basin.png",
           "Figure S36. Congo Basin NDVI comparison: 2019, 2023, difference.")

add_figure(OUTPUT_DIR / "et_comparison_maraba.png",
           "Figure S37. Maraba subregion ET comparison (MODIS MOD16A3GF) at "
           "0.05-degree resolution.")

add_figure(OUTPUT_DIR / "et_comparison_congo.png",
           "Figure S38. Congo Basin ET comparison.")

# =========================================================================
# S10. ET-OLR SCATTER DETAILS
# =========================================================================
add_heading("S10. ET vs Atmospheric Response Scatter Plots", level=1)

add_figure(OUTPUT_DIR / "et_olr_scatter.png",
           "Figure S39. ET vs OLR change scatter for Maraba (left, n=20) "
           "and Congo (right, n=150) at 1-degree resolution. "
           "Colour = cloud fraction change.")

add_figure(OUTPUT_DIR / "et_scatter_trio.png",
           "Figure S40. Maraba subregion: ET change vs OLR, cloud fraction, "
           "and CRE change per grid cell.")

# =========================================================================
# SAVE
# =========================================================================
output_path = PAPER_DIR / "Shahid_2025_Supplementary_Materials.docx"
doc.save(str(output_path))
print(f"Saved: {output_path}")
print(f"File size: {output_path.stat().st_size / 1024 / 1024:.1f} MB")

# Count figures
fig_count = sum(1 for p in doc.paragraphs if p.text.strip().startswith("Figure S"))
print(f"Supplementary figures: {fig_count}")
